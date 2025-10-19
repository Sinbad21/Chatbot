"""
Document Processor for ChatBotPlatform

Handles document upload, text extraction, chunking, and embedding generation.
Supports multiple file formats and integrates with the vector store.
"""

import os
import uuid
from typing import List, Dict, Tuple, Optional
from pathlib import Path
import logging
from openai import OpenAI
from sqlalchemy.ext.asyncio import AsyncSession
from ..core.config import settings
from ..models.user import Document, DocumentChunk
from ..core.database import get_db

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Processes documents for RAG ingestion"""

    def __init__(self, openai_client: Optional[OpenAI] = None):
        self.openai_client = openai_client or OpenAI(api_key=settings.OPENAI_API_KEY)
        self.embedding_model = "text-embedding-3-small"
        self.chunk_size = 1000  # characters
        self.chunk_overlap = 200  # characters
        self.supported_formats = ['.txt', '.pdf', '.md', '.docx']

    async def process_document(
        self,
        file_path: str,
        filename: str,
        bot_id: int,
        user_id: int,
        db: AsyncSession
    ) -> Tuple[Document, List[DocumentChunk]]:
        """
        Process a document: extract text, create chunks, generate embeddings

        Args:
            file_path: Path to the uploaded file
            filename: Original filename
            bot_id: ID of the bot
            user_id: ID of the user
            db: Database session

        Returns:
            Tuple of (Document, DocumentChunks)
        """
        try:
            # Extract text from file
            text = await self._extract_text(file_path, filename)

            # Create document record
            document = Document(
                filename=filename,
                file_path=file_path,
                bot_id=bot_id,
                user_id=user_id,
                file_size=os.path.getsize(file_path),
                content_type=self._get_content_type(filename)
            )
            db.add(document)
            await db.flush()  # Get document ID

            # Create chunks
            chunks_data = self._create_chunks(text, document.id)

            # Generate embeddings for chunks
            embeddings = await self._generate_embeddings([chunk['text'] for chunk in chunks_data])

            # Create chunk records
            chunks = []
            for i, (chunk_data, embedding) in enumerate(zip(chunks_data, embeddings)):
                chunk = DocumentChunk(
                    document_id=document.id,
                    chunk_index=i,
                    content=chunk_data['text'],
                    embedding=embedding,
                    chunk_metadata=chunk_data.get('metadata', {})
                )
                chunks.append(chunk)
                db.add(chunk)

            await db.commit()
            logger.info(f"Processed document {filename} for bot {bot_id}: {len(chunks)} chunks")

            return document, chunks

        except Exception as e:
            await db.rollback()
            logger.error(f"Error processing document {filename}: {e}")
            raise

    async def _extract_text(self, file_path: str, filename: str) -> str:
        """Extract text from various file formats"""
        file_ext = Path(filename).suffix.lower()

        if file_ext == '.txt':
            return await self._extract_text_txt(file_path)
        elif file_ext == '.pdf':
            return await self._extract_text_pdf(file_path)
        elif file_ext == '.md':
            return await self._extract_text_md(file_path)
        elif file_ext == '.docx':
            return await self._extract_text_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")

    async def _extract_text_txt(self, file_path: str) -> str:
        """Extract text from TXT files"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

    async def _extract_text_pdf(self, file_path: str) -> str:
        """Extract text from PDF files"""
        try:
            import PyPDF2
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except ImportError:
            raise ImportError("PyPDF2 is required for PDF processing. Install with: pip install PyPDF2")

    async def _extract_text_md(self, file_path: str) -> str:
        """Extract text from Markdown files"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

    async def _extract_text_docx(self, file_path: str) -> str:
        """Extract text from DOCX files"""
        try:
            from docx import Document
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except ImportError:
            raise ImportError("python-docx is required for DOCX processing. Install with: pip install python-docx")

    def _create_chunks(self, text: str, document_id: int) -> List[Dict]:
        """Split text into overlapping chunks"""
        chunks = []
        start = 0

        while start < len(text):
            end = start + self.chunk_size

            # Find sentence boundary if possible
            if end < len(text):
                # Look for sentence endings within the last 100 characters
                sentence_endings = ['. ', '! ', '? ', '\n\n']
                best_end = end
                for ending in sentence_endings:
                    pos = text.rfind(ending, start, end + 100)
                    if pos != -1 and pos > best_end - 100:
                        best_end = pos + len(ending)
                        break

                end = best_end

            chunk_text = text[start:end].strip()
            if chunk_text:
                chunks.append({
                    'text': chunk_text,
                    'chunk_id': f"{document_id}_{len(chunks)}",
                    'metadata': {
                        'start_char': start,
                        'end_char': end,
                        'document_id': document_id
                    }
                })

            # Move start position with overlap
            start = end - self.chunk_overlap

        return chunks

    async def _generate_embeddings(self, texts: List[str]) -> List[List[float]]:
        """Generate embeddings for text chunks"""
        try:
            response = self.openai_client.embeddings.create(
                input=texts,
                model=self.embedding_model
            )
            return [data.embedding for data in response.data]
        except Exception as e:
            logger.error(f"Error generating embeddings: {e}")
            raise

    def _get_content_type(self, filename: str) -> str:
        """Get MIME content type from filename"""
        ext = Path(filename).suffix.lower()
        content_types = {
            '.txt': 'text/plain',
            '.pdf': 'application/pdf',
            '.md': 'text/markdown',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        }
        return content_types.get(ext, 'application/octet-stream')

    async def delete_document_chunks(self, document_id: int, db: AsyncSession):
        """Delete all chunks for a document"""
        await db.execute(
            DocumentChunk.__table__.delete().where(DocumentChunk.document_id == document_id)
        )
        await db.commit()
        logger.info(f"Deleted chunks for document {document_id}")